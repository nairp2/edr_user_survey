import base64
from io import BytesIO
import configparser
from docx import Document
from dotenv import load_dotenv
import os
import pandas as pd
from PIL import Image
import streamlit as st
import streamlit.components.v1 as components
import toml

load_dotenv()
WORD_PATH = os.getenv("WORD_PATH")
CSV_PATH = os.getenv("CSV_PATH")
SECRETS_PATH = os.getenv("SECRETS_PATH")
CONFIG_PATH = os.getenv("CONFIG_PATH")
csv_path = os.path.join(CSV_PATH, "survey_results.csv")
word_path = os.path.join(WORD_PATH, "survey_report.docx")
secrets = toml.load(os.path.join(SECRETS_PATH, "secrets.toml"))

st.set_page_config(
    page_title="EDR Data Call User Survey",
    layout="centered",
    initial_sidebar_state="collapsed"
)

buffer = BytesIO()

img_logo = Image.open(os.path.join(SECRETS_PATH, "cms_logo.png"))
img_logo.thumbnail((150, 150))
img_logo.save(buffer, format="PNG")
img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")

config = configparser.ConfigParser(allow_no_value=True, interpolation=None)
config.read(os.path.join(CONFIG_PATH, "questions.txt"))

if "is_admin" not in st.session_state:
    st.session_state["is_admin"] = False

col1, col2, col3 = st.columns([2, 2, 2])
responses = {}

def generate_word_report(csv_path=csv_path, output_path=os.path.join(WORD_PATH, "survey_report.docx")):
    df = pd.read_csv(csv_path)
    doc = Document()

    doc.add_heading("🌍 Survey Summary Report", 0)

    for i, row in df.iterrows():
        doc.add_heading(f"Response #{i+1}", level=1)
        for col in df.columns:
            doc.add_paragraph(f"{col}: {row[col]}")
        doc.add_paragraph("-" * 40)

    # Aggregated Metrics Section
    doc.add_page_break()
    doc.add_heading("📊 Aggregated Metrics", level=0)
    doc.add_paragraph(f"Total Responses: {len(df)}")

    for col in df.columns:
        if df[col].dropna().nunique() < 20:
            doc.add_heading(f"Top Selections for: {col}", level=1)

            # For multiselect, explode them
            if df[col].dropna().astype(str).str.contains(",").fillna(False).any():
                exploded = df[col].dropna().str.split(",").explode().str.strip()
                counts = exploded.value_counts()
            else:
                counts = df[col].value_counts()

            for option, count in counts.items():
                doc.add_paragraph(f"{option}: {count}")
                
    doc.save(output_path)

st.markdown(
    """
    <div style='text-align: center;'>  <!-- Correctly centering the image -->
        <img src="data:image/png;base64,{}" width="150"/>
    </div>
    <div style='display: flex; justify-content: space-between; align-items: center;'>
        <h1 style='text-align: center; padding-left: 60px;'>🌍 EDR Data Call User Survey</h1>
    </div>
    """.format(img_base64), 
    unsafe_allow_html=True
)

with st.sidebar:
    with st.expander("🔐 Admin Login", expanded=False):
        username = st.text_input("Username", label_visibility="collapsed", key="username")
        password = st.text_input("Password", type="password", label_visibility="collapsed", key="password")
        login_button = st.button("Login")

        if login_button:
            if username in secrets['admins'] and secrets['admins'][username] == password:
                st.session_state["is_admin"] = True
                st.success(f"✅ Welcome, {username}!")
            else:
                st.session_state["is_admin"] = False
                st.error("❌ Invalid credentials")


if not config.sections():
    st.warning("⚠️ No questions found in the config file.")
else:
    with st.form("survey_form"):
        for section in config.sections():
            question = config[section].get("question", "").strip()
            if "options" in config[section]:
                options = [opt.strip() for opt in config[section]["options"].split("\n") if opt.strip()]
            else:
                options = []

            if section.startswith("Dropdown"):
                options_with_placeholder = ["-- Please select an option --"] + options
                selected = st.selectbox(question, options_with_placeholder)
                responses[question] = "" if selected == "-- Please select an option --" else selected
            elif section.startswith("Checkboxes"):
                responses[question] = st.multiselect(question, options)
            elif section.startswith("Textarea"):
                responses[question] = st.text_area(question)

        submitted = st.form_submit_button("Submit")

        if submitted:
            st.success("🎉 Thank you for your submission!")
            st.write("Here’s what you submitted:")
            for q, a in responses.items():
                if isinstance(a, list):
                    st.write(f"**{q}**: {', '.join(a)}")
                else:
                    st.write(f"**{q}**: {a}")

            response = {q: ", ".join(a) if isinstance(a, list) else a for q, a in responses.items()}

            if os.path.exists(csv_path):
                df = pd.read_csv(csv_path, encoding="utf-8")
            else:
                df = pd.DataFrame(columns=response.keys())

            df = pd.concat([df, pd.DataFrame([response])], ignore_index=True)
            df.to_csv(csv_path, index=False, encoding="utf-8")

            generate_word_report(csv_path=csv_path, output_path=word_path)
            st.success("✅ Your response has been recorded!")

if st.session_state.get("is_admin"):
    with open(word_path, "rb") as f:
        st.download_button(
            label="Download Survey Report",
            data=f,
            file_name=word_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )